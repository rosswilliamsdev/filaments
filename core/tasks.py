"""
AI processing pipeline (backend-planning-doc.md → Business Logic & Edge Cases).

Celery chain: transcribe/extract → Claude extraction (summary + key ideas +
action items + tags, one call) → embedding → auto-link → finalize.

Failure policy:
- Persist-and-resume: every step persists its output on completion and skips
  itself when that output already exists, so a requeued chain resumes instead
  of restarting (never re-pay Whisper because a later step failed).
- Transient errors (rate limits, timeouts, 5xx) retry ~3x with exponential
  backoff; anything else is terminal for that step.
- Critical path (transcribe/extract) terminal failure → 'failed', chain stops.
  Enrichment terminal failure (Claude, embedding, auto-link) → log and
  continue: the filament still reaches 'done', degraded, and a later re-run
  backfills the missing enrichment.
- A URL that fetches but yields no extractable text (paywall, login wall,
  JS-rendered page) fails identically on every attempt — non-retryable.
"""

import io
import json
import logging
import re

import anthropic
import docx
import openai
import pymupdf
import requests
import trafilatura
from pgvector.django import CosineDistance
from celery import chain, shared_task
from celery.exceptions import MaxRetriesExceededError
from django.conf import settings
from django.db import transaction
from django.utils import timezone

from .models import ActionItem, Filament, FilamentLink, Tag
from .s3 import delete_object, download_object

logger = logging.getLogger(__name__)

CLAUDE_MODEL = "claude-sonnet-4-6"
WHISPER_MODEL = "whisper-1"
EMBEDDING_MODEL = "text-embedding-3-small"

# Auto-link decisions (PRD open Q #3, settled 2026-06): minimum cosine
# similarity to create a link, and at most this many links per new filament.
AUTO_LINK_THRESHOLD = 0.75
AUTO_LINK_LIMIT = 5

MAX_KEY_IDEAS = 5
MAX_TAGS = 8
TAG_MAX_LENGTH = 64  # Tag.name max_length

EXTRACTION_MAX_CHARS = 150_000  # bound prompt cost on very large documents
EMBED_MAX_CHARS = 30_000  # ~8k tokens, text-embedding-3-small's input limit

TRANSIENT_ERRORS = (
    anthropic.APIConnectionError,
    anthropic.RateLimitError,
    anthropic.InternalServerError,
    openai.APIConnectionError,
    openai.RateLimitError,
    openai.InternalServerError,
    requests.ConnectionError,
    requests.Timeout,
)

RETRY_POLICY = {"max_retries": 3, "retry_backoff": True, "retry_jitter": True}


class NonRetryableError(Exception):
    """Fails identically on every attempt — fail fast instead of retrying."""


class MalformedExtraction(Exception):
    """Claude output that doesn't match the extraction schema."""


def get_anthropic_client():
    return anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)


def get_openai_client():
    return openai.OpenAI(api_key=settings.OPENAI_API_KEY)


# ---------------------------------------------------------------------------
# Chain plumbing
# ---------------------------------------------------------------------------


def build_pipeline(filament_id: str):
    return chain(
        transcribe_extract.si(filament_id),
        enrich.si(filament_id),
        generate_embedding.si(filament_id),
        auto_link.si(filament_id),
        finalize.si(filament_id),
    )


@shared_task
def process_filament(filament_id: str) -> None:
    """Pipeline entry point, enqueued by the status-gated POST /filaments/{id}/process."""
    build_pipeline(filament_id).apply_async()


def _get_filament(filament_id):
    """Load the filament for a pipeline step; None means skip the step."""
    filament = Filament.objects.filter(pk=filament_id, deleted_at__isnull=True).first()
    if filament is None or filament.status == Filament.Status.FAILED:
        return None
    return filament


def _mark_failed(filament_id):
    Filament.objects.filter(
        pk=filament_id, status=Filament.Status.PROCESSING
    ).update(status=Filament.Status.FAILED, updated_at=timezone.now())


# ---------------------------------------------------------------------------
# Step 1 — transcribe (voice) / extract (document). Critical path.
# ---------------------------------------------------------------------------


@shared_task(bind=True, **RETRY_POLICY)
def transcribe_extract(self, filament_id: str) -> None:
    """
    Populate `body` (+ `transcript` for voice). Without this there is no
    readable filament, so terminal failure marks 'failed' and stops the chain.
    """
    filament = _get_filament(filament_id)
    if filament is None:
        return
    try:
        if filament.type == Filament.Type.VOICE:
            _transcribe_voice(filament)
        elif filament.type == Filament.Type.DOCUMENT:
            _extract_document(filament)
        # Text notes already carry their body from capture.
    except TRANSIENT_ERRORS as exc:
        try:
            raise self.retry(exc=exc)
        except MaxRetriesExceededError:
            _mark_failed(filament_id)
            raise exc
    except Exception:
        _mark_failed(filament_id)
        raise


def _transcribe_voice(filament):
    if filament.transcript is not None and filament.body:
        return  # persist-and-resume: never re-call Whisper
    if not filament.source_key:
        raise NonRetryableError("voice filament has no source_key")

    source_key = filament.source_key
    audio = download_object(source_key)
    response = get_openai_client().audio.transcriptions.create(
        model=WHISPER_MODEL,
        file=(source_key.rsplit("/", 1)[-1], audio),
        response_format="verbose_json",
    )
    # `speaker` stays null in v1: whisper-1 has no diarization and
    # multi-speaker capture is a non-goal (backend doc → speaker labels).
    filament.transcript = [
        {"start": segment.start, "end": segment.end, "speaker": None, "text": segment.text}
        for segment in (response.segments or [])
    ]
    filament.body = (response.text or "").strip()
    filament.save(update_fields=["transcript", "body", "updated_at"])

    _discard_audio(filament, source_key)


def _discard_audio(filament, source_key):
    """
    Voice audio is transient — we keep the transcript, not the recording. Once
    the transcript is persisted the guard in `_transcribe_voice` blocks any
    re-transcription, so the S3 object is never needed again and we delete it now
    instead of storing it.

    Best-effort: a delete failure leaves the object in place (the soft-delete
    sweep still reaps it when the filament is deleted), so we only null
    `source_key` once the object is actually gone — keeping a live key pointing
    at a live object for the sweep to find.
    """
    if not settings.USE_S3:
        return
    try:
        delete_object(source_key)
    except Exception:
        logger.warning(
            "could not delete transcribed audio %s; leaving for the sweep", source_key
        )
        return
    filament.source_key = None
    filament.save(update_fields=["source_key", "updated_at"])


def _extract_document(filament):
    if filament.body:
        return  # persist-and-resume
    source = filament.source_key or ""
    if not source:
        raise NonRetryableError("document filament has no source_key")

    # URL captures store the URL in source_key (no S3 object); uploaded files
    # store an S3 key ending in their format extension. The capture API for
    # URLs isn't built yet — this dispatch is forward-compatible with it.
    if source.startswith(("http://", "https://")):
        text = _extract_url(source)
    else:
        text = _extract_file(source, download_object(source))

    if not text.strip():
        raise NonRetryableError("no extractable text in document")
    filament.body = text.strip()
    filament.save(update_fields=["body", "updated_at"])


def _extract_file(source_key: str, data: bytes) -> str:
    """Dispatch on the upload's extension (set at create time in s3.py)."""
    ext = source_key.rsplit(".", 1)[-1].lower() if "." in source_key else ""
    if ext == "pdf":
        return _extract_pdf(data)
    if ext == "docx":
        return _extract_docx(data)
    # md / markdown / txt are already plain text — decode and pass through.
    return data.decode("utf-8", errors="replace")


def _extract_url(url: str) -> str:
    response = requests.get(url, timeout=30)
    if response.status_code >= 500:
        raise requests.ConnectionError(f"{url} returned {response.status_code}")
    if not response.ok:
        raise NonRetryableError(f"{url} returned {response.status_code}")
    text = trafilatura.extract(response.text)
    if not text or not text.strip():
        # Paywall/login/JS page: retrying wastes time (backend doc → retry policy).
        raise NonRetryableError("fetched URL but no extractable text")
    return text


def _extract_pdf(data: bytes) -> str:
    with pymupdf.open(stream=data, filetype="pdf") as doc:
        return "\n".join(page.get_text() for page in doc)


def _extract_docx(data: bytes) -> str:
    # Paragraphs carry the prose; table cells often hold real content too, so
    # append them (tab-joined per row) rather than dropping them silently.
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Step 2 — Claude extraction (enrichment). One call for summary, key ideas,
# action items, and tags (PRD open Q #4, settled 2026-06).
# ---------------------------------------------------------------------------

# Structured outputs don't support maxItems, so the ≤5 / ≤8 caps are enforced
# in _clean_extraction below.
EXTRACTION_SCHEMA = {
    "type": "object",
    "properties": {
        "summary": {"type": "string"},
        "key_ideas": {"type": "array", "items": {"type": "string"}},
        "action_items": {"type": "array", "items": {"type": "string"}},
        "tags": {"type": "array", "items": {"type": "string"}},
    },
    "required": ["summary", "key_ideas", "action_items", "tags"],
    "additionalProperties": False,
}

# One prompt for all three input types; only the framing line changes, the
# output schema is identical (settled decision). The prose below is the
# tunable part — what makes a good summary/tag set for Ross's thinking style.
EXTRACTION_FRAMING = {
    Filament.Type.VOICE: (
        "The text below is a transcribed voice memo — spoken thought, so expect "
        "filler words and meandering structure. Capture the intent, not the "
        "verbatim phrasing."
    ),
    Filament.Type.DOCUMENT: (
        "The text below was extracted from a saved document or article. Capture "
        "the author's argument and anything worth keeping for later."
    ),
    Filament.Type.TEXT: (
        "The text below is a written note — already deliberate, so stay close "
        "to its own wording."
    ),
}

EXTRACTION_PROMPT = """\
You are the extraction step in Filaments, a personal knowledge graph. {framing}

<filament title="{title}">
{body}
</filament>

Produce:
- summary: 2-3 sentences on what this filament is about and why it mattered \
enough to capture. Write it to be useful when skimming a timeline months later.
- key_ideas: at most {max_key_ideas} distinct ideas worth remembering, each a \
single self-contained sentence.
- action_items: only concrete, doable to-dos the author committed to or clearly \
implied. An empty list is the right answer when there are none.
- tags: at most {max_tags} reusable topic tags — lowercase, hyphenated nouns \
like "deep-work". Prefer broad recurring themes over one-off specifics.
"""


@shared_task(bind=True, **RETRY_POLICY)
def enrich(self, filament_id: str) -> None:
    """
    Claude extraction → summary, key_ideas, ActionItem rows, Tag M2M.
    Enrichment: terminal failure logs and returns so the chain continues.
    """
    filament = _get_filament(filament_id)
    if filament is None or filament.summary or not filament.body.strip():
        return
    try:
        extraction = _run_extraction(filament)
    except TRANSIENT_ERRORS as exc:
        try:
            raise self.retry(exc=exc)
        except MaxRetriesExceededError:
            logger.warning("enrich(%s): retries exhausted — continuing degraded", filament_id)
            return
    except MalformedExtraction:
        logger.warning(
            "enrich(%s): malformed output after one retry — continuing degraded", filament_id
        )
        return
    except Exception:
        logger.exception("enrich(%s): failed — continuing degraded", filament_id)
        return
    _persist_extraction(filament, extraction)


def _run_extraction(filament) -> dict:
    client = get_anthropic_client()
    prompt = EXTRACTION_PROMPT.format(
        framing=EXTRACTION_FRAMING[filament.type],
        title=filament.title or "(untitled)",
        body=filament.body[:EXTRACTION_MAX_CHARS],
        max_key_ideas=MAX_KEY_IDEAS,
        max_tags=MAX_TAGS,
    )
    last_error = None
    for _ in range(2):  # malformed → retry once → degraded (settled decision)
        response = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=2048,
            messages=[{"role": "user", "content": prompt}],
            output_config={"format": {"type": "json_schema", "schema": EXTRACTION_SCHEMA}},
        )
        try:
            return _clean_extraction(_parse_extraction(response))
        except MalformedExtraction as exc:
            last_error = exc
    raise last_error


def _parse_extraction(response) -> dict:
    text = next((block.text for block in response.content if block.type == "text"), "")
    try:
        return json.loads(text)
    except (TypeError, ValueError) as exc:
        raise MalformedExtraction(f"not valid JSON: {exc}") from exc


def _clean_extraction(data) -> dict:
    """Validate before persisting; raises MalformedExtraction on schema drift."""
    if not isinstance(data, dict):
        raise MalformedExtraction("not an object")
    summary = data.get("summary")
    if not isinstance(summary, str) or not summary.strip():
        raise MalformedExtraction("summary missing or empty")

    def strings(value):
        if not isinstance(value, list):
            raise MalformedExtraction("expected a list of strings")
        return [item.strip() for item in value if isinstance(item, str) and item.strip()]

    tags = []
    for name in strings(data.get("tags")):
        slug = re.sub(r"[\s_]+", "-", name.lower()).strip("-")[:TAG_MAX_LENGTH]
        if slug and slug not in tags:
            tags.append(slug)

    return {
        "summary": summary.strip(),
        "key_ideas": strings(data.get("key_ideas"))[:MAX_KEY_IDEAS],
        "action_items": strings(data.get("action_items")),
        "tags": tags[:MAX_TAGS],
    }


def _persist_extraction(filament, extraction):
    # Atomic so a crash mid-persist leaves summary empty and the step re-runs
    # cleanly (the summary guard is what makes this step idempotent).
    with transaction.atomic():
        filament.summary = extraction["summary"]
        filament.key_ideas = extraction["key_ideas"]
        filament.save(update_fields=["summary", "key_ideas", "updated_at"])
        ActionItem.objects.bulk_create(
            ActionItem(filament=filament, text=text) for text in extraction["action_items"]
        )
        filament.tags.add(
            *(Tag.objects.get_or_create(name=name)[0] for name in extraction["tags"])
        )


# ---------------------------------------------------------------------------
# Step 3 — embedding (enrichment)
# ---------------------------------------------------------------------------


@shared_task(bind=True, **RETRY_POLICY)
def generate_embedding(self, filament_id: str) -> None:
    """Embed from `body` only — never from transcript or summary."""
    filament = _get_filament(filament_id)
    if filament is None or filament.embedding is not None or not filament.body.strip():
        return
    try:
        response = get_openai_client().embeddings.create(
            model=EMBEDDING_MODEL, input=filament.body[:EMBED_MAX_CHARS]
        )
    except TRANSIENT_ERRORS as exc:
        try:
            raise self.retry(exc=exc)
        except MaxRetriesExceededError:
            logger.warning(
                "generate_embedding(%s): retries exhausted — continuing degraded", filament_id
            )
            return
    except Exception:
        logger.exception("generate_embedding(%s): failed — continuing degraded", filament_id)
        return
    filament.embedding = response.data[0].embedding
    filament.save(update_fields=["embedding", "updated_at"])


# ---------------------------------------------------------------------------
# Step 4 — auto-link (enrichment)
# ---------------------------------------------------------------------------


@shared_task(bind=True, **RETRY_POLICY)
def auto_link(self, filament_id: str) -> None:
    """
    pgvector cosine similarity against existing filaments. All link creation
    goes through FilamentLink.create_link() (canonical ordering, upsert), so
    re-processing refreshes scores instead of duplicating rows.
    """
    filament = _get_filament(filament_id)
    if filament is None:
        return
    if filament.embedding is None:
        logger.info("auto_link(%s): no embedding (degraded) — skipping", filament_id)
        return
    try:
        candidates = (
            Filament.objects.filter(deleted_at__isnull=True, embedding__isnull=False)
            .exclude(pk=filament.pk)
            .annotate(distance=CosineDistance("embedding", filament.embedding))
            .filter(distance__lte=1 - AUTO_LINK_THRESHOLD)
            .order_by("distance")[:AUTO_LINK_LIMIT]
        )
        for other in candidates:
            FilamentLink.create_link(filament, other, score=1 - other.distance)
    except Exception:
        logger.exception("auto_link(%s): failed — continuing degraded", filament_id)


# ---------------------------------------------------------------------------
# Step 5 — finalize
# ---------------------------------------------------------------------------


@shared_task
def finalize(filament_id: str) -> None:
    """
    Mark done. Conditional on 'processing' so a stale chain delivery can never
    resurrect a filament the critical step already marked 'failed'.
    """
    updated = Filament.objects.filter(
        pk=filament_id, deleted_at__isnull=True, status=Filament.Status.PROCESSING
    ).update(status=Filament.Status.DONE, updated_at=timezone.now())
    if updated:
        logger.info("process_filament(%s): done", filament_id)
