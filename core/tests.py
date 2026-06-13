import json
import math
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

from django.contrib.auth.models import User
from django.db import IntegrityError
from django.db.models import Q
from django.test import TestCase, override_settings
from django.utils import timezone
from rest_framework.test import APITestCase

from .models import ActionItem, Filament, FilamentLink, Tag
from .tasks import AUTO_LINK_LIMIT, AUTO_LINK_THRESHOLD, build_pipeline


class AuthedAPITestCase(APITestCase):
    @classmethod
    def setUpTestData(cls):
        cls.user = User.objects.create_user(
            username="test@example.com", email="test@example.com"
        )

    def setUp(self):
        self.client.force_authenticate(self.user)


class AuthBoundaryTests(APITestCase):
    def test_unauthenticated_requests_are_rejected(self):
        for url in ("/api/v1/filaments", "/api/v1/search?q=x", "/api/v1/tags"):
            self.assertEqual(self.client.get(url).status_code, 401, url)


class FilamentAPITests(AuthedAPITestCase):
    def test_text_note_create_and_process_handshake(self):
        res = self.client.post(
            "/api/v1/filaments",
            {"type": "text", "title": "Note", "body": "Hello world"},
            format="json",
        )
        self.assertEqual(res.status_code, 201)
        self.assertIsNone(res.data["upload_url"])
        fid = res.data["filament_id"]

        with patch("core.views.process_filament") as task:
            res = self.client.post(f"/api/v1/filaments/{fid}/process")
            self.assertEqual(res.status_code, 202)
            self.assertTrue(res.data["enqueued"])
            self.assertEqual(res.data["status"], "processing")
            task.delay.assert_called_once_with(fid)

            # Duplicate call (retry/double-tap): no-op success, no second chain
            res = self.client.post(f"/api/v1/filaments/{fid}/process")
            self.assertEqual(res.status_code, 202)
            self.assertFalse(res.data["enqueued"])
            task.delay.assert_called_once()

        self.assertEqual(Filament.objects.get(pk=fid).pipeline_attempts, 1)

    def test_text_note_requires_body(self):
        res = self.client.post("/api/v1/filaments", {"type": "text"}, format="json")
        self.assertEqual(res.status_code, 400)

    @override_settings(USE_S3=False)  # don't depend on the dev .env having S3
    def test_file_create_without_s3_fails_loudly_and_leaves_no_row(self):
        res = self.client.post("/api/v1/filaments", {"type": "voice"}, format="json")
        self.assertEqual(res.status_code, 503)
        self.assertEqual(Filament.objects.count(), 0)

    def test_list_excludes_archived_and_deleted_by_default(self):
        Filament.objects.create(type="text", title="live", body="x")
        Filament.objects.create(type="text", title="archived", body="x", archived=True)
        Filament.objects.create(
            type="text", title="deleted", body="x", deleted_at=timezone.now()
        )

        res = self.client.get("/api/v1/filaments")
        self.assertEqual([r["title"] for r in res.data["results"]], ["live"])

        res = self.client.get("/api/v1/filaments?archived=true")
        self.assertEqual([r["title"] for r in res.data["results"]], ["archived"])

    def test_patch_title_and_tags(self):
        f = Filament.objects.create(type="text", title="old", body="b")
        res = self.client.patch(
            f"/api/v1/filaments/{f.id}",
            {"title": "new", "tags": ["focus", "deep-work", "focus"]},
            format="json",
        )
        self.assertEqual(res.status_code, 200)
        self.assertEqual(res.data["title"], "new")
        self.assertEqual(sorted(res.data["tags"]), ["deep-work", "focus"])
        self.assertEqual(Tag.objects.count(), 2)

    def test_pipeline_fields_are_read_only(self):
        f = Filament.objects.create(type="text", title="t", body="b")
        res = self.client.patch(
            f"/api/v1/filaments/{f.id}", {"status": "done"}, format="json"
        )
        self.assertEqual(res.status_code, 200)
        f.refresh_from_db()
        self.assertEqual(f.status, Filament.Status.PENDING_UPLOAD)

    def test_soft_delete_tombstones_and_hides(self):
        f = Filament.objects.create(type="text", title="t", body="b")
        res = self.client.delete(f"/api/v1/filaments/{f.id}")
        self.assertEqual(res.status_code, 204)
        f.refresh_from_db()
        self.assertIsNotNone(f.deleted_at)  # row survives for the sweep
        self.assertEqual(self.client.get(f"/api/v1/filaments/{f.id}").status_code, 404)

    def test_action_item_toggle(self):
        f = Filament.objects.create(type="text", title="t", body="b")
        item = ActionItem.objects.create(filament=f, text="do the thing")
        res = self.client.patch(
            f"/api/v1/filaments/{f.id}/action-items/{item.id}",
            {"done": True},
            format="json",
        )
        self.assertEqual(res.status_code, 200)
        item.refresh_from_db()
        self.assertTrue(item.done)


class DocumentUploadTests(AuthedAPITestCase):
    """Bulk web upload supports PDF, Word, and markdown. Each file's format is
    carried by `filename` → validated → baked into the S3 key extension, which
    the pipeline later dispatches on (see tasks._extract_file)."""

    @patch("core.views.generate_upload_url", return_value="https://s3.example/put")
    def test_create_uses_filename_extension(self, _url):
        for filename, expected_ext in [
            ("Report.pdf", ".pdf"),
            ("Notes.DOCX", ".docx"),  # extension match is case-insensitive
            ("readme.md", ".md"),
            ("scratch.txt", ".txt"),
        ]:
            res = self.client.post(
                "/api/v1/filaments",
                {"type": "document", "filename": filename},
                format="json",
            )
            self.assertEqual(res.status_code, 201, filename)
            self.assertEqual(res.data["upload_url"], "https://s3.example/put")
            f = Filament.objects.get(pk=res.data["filament_id"])
            self.assertTrue(f.source_key.endswith(expected_ext), f.source_key)

    def test_create_rejects_unsupported_type_and_strands_no_row(self):
        res = self.client.post(
            "/api/v1/filaments",
            {"type": "document", "filename": "malware.exe"},
            format="json",
        )
        self.assertEqual(res.status_code, 400)
        self.assertEqual(Filament.objects.count(), 0)

    def test_create_rejects_document_without_filename(self):
        # No filename → no known format → reject before minting an upload URL.
        res = self.client.post(
            "/api/v1/filaments", {"type": "document"}, format="json"
        )
        self.assertEqual(res.status_code, 400)
        self.assertEqual(Filament.objects.count(), 0)


class DocumentExtractionTests(TestCase):
    """The pipeline dispatches extraction on the source key's extension."""

    def test_dispatch_routes_by_extension(self):
        from .tasks import _extract_file

        with patch("core.tasks._extract_pdf", return_value="pdf text") as pdf:
            self.assertEqual(_extract_file("document/x.pdf", b"%PDF"), "pdf text")
        pdf.assert_called_once()

        with patch("core.tasks._extract_docx", return_value="word text") as docx_:
            self.assertEqual(_extract_file("document/x.docx", b"PK\x03\x04"), "word text")
        docx_.assert_called_once()

    def test_markdown_passes_through_as_plain_text(self):
        from .tasks import _extract_file

        body = "# Heading\n\nSome **markdown** body."
        self.assertEqual(_extract_file("document/x.md", body.encode()), body)

    def test_docx_extraction_reads_paragraphs_and_tables(self):
        import io

        import docx

        from .tasks import _extract_docx

        document = docx.Document()
        document.add_paragraph("Opening paragraph.")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "left cell"
        table.rows[0].cells[1].text = "right cell"
        buf = io.BytesIO()
        document.save(buf)

        out = _extract_docx(buf.getvalue())
        self.assertIn("Opening paragraph.", out)
        self.assertIn("left cell", out)
        self.assertIn("right cell", out)


class SearchTests(AuthedAPITestCase):
    def test_fts_matches_and_filters(self):
        Filament.objects.create(
            type="text", title="Deep work", body="attention and focus"
        )
        Filament.objects.create(type="voice", title="Groceries", body="milk eggs")

        res = self.client.get("/api/v1/search?q=attention")
        self.assertEqual([r["title"] for r in res.data["results"]], ["Deep work"])

        res = self.client.get("/api/v1/search?q=attention&type=voice")
        self.assertEqual(res.data["results"], [])

        res = self.client.get("/api/v1/search")  # no q → empty, not an error
        self.assertEqual(res.data["results"], [])

    def test_search_excludes_soft_deleted(self):
        Filament.objects.create(
            type="text", title="gone", body="attention", deleted_at=timezone.now()
        )
        res = self.client.get("/api/v1/search?q=attention")
        self.assertEqual(res.data["results"], [])


class TagListTests(AuthedAPITestCase):
    def test_counts_live_filaments_only(self):
        tag = Tag.objects.create(name="focus")
        live = Filament.objects.create(type="text", title="x", body="y")
        dead = Filament.objects.create(
            type="text", title="z", body="y", deleted_at=timezone.now()
        )
        live.tags.add(tag)
        dead.tags.add(tag)

        res = self.client.get("/api/v1/tags")
        self.assertEqual(res.data, [{"id": tag.id, "name": "focus", "count": 1}])


class FilamentLinkTests(AuthedAPITestCase):
    def test_create_link_canonicalizes_and_refreshes_score(self):
        a = Filament.objects.create(type="text", title="a", body="x")
        b = Filament.objects.create(type="text", title="b", body="y")

        first = FilamentLink.create_link(a, b, 0.9)
        second = FilamentLink.create_link(b, a, 0.5)  # reversed re-process

        self.assertEqual(first.pk, second.pk)
        link = FilamentLink.objects.get()
        self.assertEqual(link.score, 0.5)
        self.assertLess(link.source_id, link.target_id)

    def test_reversed_direct_insert_rejected_by_db(self):
        a = Filament.objects.create(type="text", title="a", body="x")
        b = Filament.objects.create(type="text", title="b", body="y")
        lo, hi = sorted([a, b], key=lambda f: f.pk)
        with self.assertRaises(IntegrityError):
            FilamentLink.objects.create(source=hi, target=lo, score=0.5)

    def test_self_link_rejected(self):
        a = Filament.objects.create(type="text", title="a", body="x")
        with self.assertRaises(ValueError):
            FilamentLink.create_link(a, a, 1.0)

    def test_links_render_undirected_in_detail(self):
        a = Filament.objects.create(type="text", title="a", body="x")
        b = Filament.objects.create(type="text", title="b", body="y")
        FilamentLink.create_link(a, b, 0.9)

        for filament, other in ((a, b), (b, a)):
            res = self.client.get(f"/api/v1/filaments/{filament.id}")
            self.assertEqual(
                [link["filament_id"] for link in res.data["links"]], [str(other.id)]
            )


# ---------------------------------------------------------------------------
# AI pipeline integration tests (mocked Whisper/Claude/OpenAI clients).
# The chain runs eagerly via chain.apply(); a failing task's exception
# propagates when the chain prepares the next link, so "chain stopped" is
# directly observable.
# ---------------------------------------------------------------------------


def vec(x, y):
    """A unit 1536-dim vector in the (e0, e1) plane: cosine sim vs vec(1,0) == x."""
    v = [0.0] * 1536
    v[0], v[1] = x, y
    return v


def sim_vec(similarity):
    return vec(similarity, math.sqrt(1 - similarity**2))


def claude_response(payload):
    """Structured-output response: one text block of schema-valid JSON."""
    return SimpleNamespace(content=[SimpleNamespace(type="text", text=json.dumps(payload))])


def embedding_response(vector):
    return SimpleNamespace(data=[SimpleNamespace(embedding=vector)])


def whisper_response(text, segments):
    return SimpleNamespace(text=text, segments=[SimpleNamespace(**s) for s in segments])


EXTRACTION = {
    "summary": "A reflection on attention and deep work.",
    "key_ideas": ["one", "two", "three", "four", "five", "six"],  # 6 → clipped to 5
    "action_items": ["Block two hours tomorrow", "Re-read chapter 3"],
    "tags": ["Deep Work", "focus", "FOCUS", "attention economics"],  # → 3 normalized
}


class PipelineIntegrationTests(TestCase):
    def _clients(self, anthropic_client, openai_client):
        return (
            patch("core.tasks.get_anthropic_client", return_value=anthropic_client),
            patch("core.tasks.get_openai_client", return_value=openai_client),
            patch("core.tasks.download_object", return_value=b"raw-bytes"),
        )

    def test_text_filament_full_chain(self):
        filament = Filament.objects.create(
            type="text", title="Note", body="Thoughts on attention",
            status=Filament.Status.PROCESSING,
        )
        sims = [0.99, 0.95, 0.90, 0.85, 0.80, 0.78]  # 6 above threshold → top 5 kept
        candidates = {
            s: Filament.objects.create(
                type="text", title=f"c{s}", body="x", status="done", embedding=sim_vec(s)
            )
            for s in sims
        }
        Filament.objects.create(  # below threshold → never linked
            type="text", title="far", body="x", status="done", embedding=sim_vec(0.70)
        )
        Filament.objects.create(  # soft-deleted → excluded even at sim 0.97
            type="text", title="gone", body="x", status="done",
            embedding=sim_vec(0.97), deleted_at=timezone.now(),
        )
        Filament.objects.create(  # no embedding → excluded
            type="text", title="raw", body="x", status="processing"
        )

        anthropic_client = MagicMock()
        anthropic_client.messages.create.return_value = claude_response(EXTRACTION)
        openai_client = MagicMock()
        openai_client.embeddings.create.return_value = embedding_response(vec(1.0, 0.0))

        p1, p2, p3 = self._clients(anthropic_client, openai_client)
        with p1, p2, p3:
            build_pipeline(str(filament.id)).apply()

        filament.refresh_from_db()
        self.assertEqual(filament.status, Filament.Status.DONE)
        self.assertEqual(filament.summary, EXTRACTION["summary"])
        self.assertEqual(filament.key_ideas, ["one", "two", "three", "four", "five"])
        self.assertEqual(
            [item.text for item in filament.action_items.order_by("id")],
            EXTRACTION["action_items"],
        )
        self.assertEqual(
            set(filament.tags.values_list("name", flat=True)),
            {"deep-work", "focus", "attention-economics"},
        )
        self.assertIsNotNone(filament.embedding)
        openai_client.audio.transcriptions.create.assert_not_called()  # text → no Whisper

        links = FilamentLink.objects.filter(Q(source=filament) | Q(target=filament))
        self.assertEqual(links.count(), AUTO_LINK_LIMIT)
        partners = {
            link.target_id if link.source_id == filament.id else link.source_id
            for link in links
        }
        self.assertEqual(partners, {candidates[s].id for s in sims[:AUTO_LINK_LIMIT]})
        for link in links:
            self.assertGreaterEqual(link.score, AUTO_LINK_THRESHOLD - 1e-6)
            self.assertLess(link.source_id, link.target_id)  # via create_link()

    def test_embedding_failure_degrades_then_resume_skips_completed_steps(self):
        filament = Filament.objects.create(
            type="voice", title="Memo", status=Filament.Status.PROCESSING,
            source_key="voice/abc.m4a",
        )
        anthropic_client = MagicMock()
        anthropic_client.messages.create.return_value = claude_response(EXTRACTION)
        openai_client = MagicMock()
        openai_client.audio.transcriptions.create.return_value = whisper_response(
            "hello world", [{"start": 0.0, "end": 2.5, "text": "hello world"}]
        )
        openai_client.embeddings.create.side_effect = RuntimeError("embedding service down")

        p1, p2, p3 = self._clients(anthropic_client, openai_client)
        with p1, p2, p3:
            build_pipeline(str(filament.id)).apply()

        filament.refresh_from_db()
        self.assertEqual(filament.status, Filament.Status.DONE)  # degraded, not blocked
        self.assertEqual(filament.body, "hello world")
        self.assertEqual(
            filament.transcript,
            [{"start": 0.0, "end": 2.5, "speaker": None, "text": "hello world"}],
        )
        self.assertEqual(filament.summary, EXTRACTION["summary"])
        self.assertIsNone(filament.embedding)
        self.assertEqual(FilamentLink.objects.count(), 0)  # auto-link skipped
        self.assertEqual(openai_client.audio.transcriptions.create.call_count, 1)

        # Re-run the chain (stuck-sweep / manual-retry semantics): completed
        # steps must not re-execute — the expensive calls aren't re-paid.
        filament.status = Filament.Status.PROCESSING
        filament.save(update_fields=["status"])
        p1, p2, p3 = self._clients(anthropic_client, openai_client)
        with p1, p2, p3:
            build_pipeline(str(filament.id)).apply()

        filament.refresh_from_db()
        self.assertEqual(filament.status, Filament.Status.DONE)
        self.assertEqual(openai_client.audio.transcriptions.create.call_count, 1)  # no re-Whisper
        self.assertEqual(anthropic_client.messages.create.call_count, 1)  # summary persisted
        self.assertEqual(openai_client.embeddings.create.call_count, 2)  # backfill attempted
        self.assertIsNone(filament.embedding)  # still degraded

    def test_transcription_failure_marks_failed_and_stops_chain(self):
        filament = Filament.objects.create(
            type="voice", title="Memo", status=Filament.Status.PROCESSING,
            source_key="voice/abc.m4a",
        )
        anthropic_client = MagicMock()
        openai_client = MagicMock()
        openai_client.audio.transcriptions.create.side_effect = RuntimeError("whisper down")

        p1, p2, p3 = self._clients(anthropic_client, openai_client)
        with p1, p2, p3:
            with self.assertRaises(RuntimeError):
                build_pipeline(str(filament.id)).apply()

        filament.refresh_from_db()
        self.assertEqual(filament.status, Filament.Status.FAILED)
        self.assertEqual(filament.summary, "")
        anthropic_client.messages.create.assert_not_called()  # chain stopped
        openai_client.embeddings.create.assert_not_called()

    def test_malformed_extraction_retries_once_then_degrades(self):
        filament = Filament.objects.create(
            type="text", title="Note", body="Some text",
            status=Filament.Status.PROCESSING,
        )
        anthropic_client = MagicMock()
        anthropic_client.messages.create.return_value = claude_response({"summary": ""})
        openai_client = MagicMock()
        openai_client.embeddings.create.return_value = embedding_response(vec(1.0, 0.0))

        p1, p2, p3 = self._clients(anthropic_client, openai_client)
        with p1, p2, p3:
            build_pipeline(str(filament.id)).apply()

        filament.refresh_from_db()
        self.assertEqual(filament.status, Filament.Status.DONE)
        self.assertEqual(anthropic_client.messages.create.call_count, 2)  # retry once, then degrade
        self.assertEqual(filament.summary, "")
        self.assertEqual(filament.tags.count(), 0)
        self.assertEqual(filament.action_items.count(), 0)
        self.assertIsNotNone(filament.embedding)  # embeds from body, unaffected by enrichment


# ---------------------------------------------------------------------------
# /ask (RAG) — mocked OpenAI embedding + Claude clients
# ---------------------------------------------------------------------------


class AskTests(AuthedAPITestCase):
    def _ask(self, question, anthropic_client, openai_client):
        with patch("core.rag.get_anthropic_client", return_value=anthropic_client), \
             patch("core.rag.get_openai_client", return_value=openai_client):
            return self.client.post("/api/v1/ask", {"question": question}, format="json")

    def test_structured_answer_with_cited_sources(self):
        near = Filament.objects.create(
            type="document", title="Deep Work Refined", body="attention is finite",
            summary="On focus.", status="done", embedding=sim_vec(0.99),
        )
        far = Filament.objects.create(
            type="text", title="Groceries", body="milk", status="done",
            embedding=sim_vec(0.10),
        )
        Filament.objects.create(  # deleted → never retrieved
            type="text", title="gone", body="x", status="done",
            embedding=sim_vec(0.98), deleted_at=timezone.now(),
        )

        anthropic_client = MagicMock()
        anthropic_client.messages.create.return_value = claude_response({
            "answer": [
                {"text": "A recurring theme is ", "citation": None},
                {"text": "finite attention", "citation": 1},
                {"text": " with an out-of-range citation here", "citation": 99},
            ],
            "follow_ups": ["What did I say about focus rituals?"],
        })
        openai_client = MagicMock()
        openai_client.embeddings.create.return_value = embedding_response(vec(1.0, 0.0))

        res = self._ask("what themes recur?", anthropic_client, openai_client)
        self.assertEqual(res.status_code, 200)
        self.assertEqual(res.data["answer"][1]["citation"], 1)
        self.assertIsNone(res.data["answer"][2]["citation"])  # 99 degraded to prose
        self.assertEqual(res.data["follow_ups"], ["What did I say about focus rituals?"])
        # Only cited sources are returned; citation 1 = nearest filament
        self.assertEqual(len(res.data["sources"]), 1)
        source = res.data["sources"][0]
        self.assertEqual(source["citation"], 1)
        self.assertEqual(source["filament_id"], str(near.id))
        self.assertEqual(source["type"], "document")
        self.assertEqual(source["snippet"], "On focus.")  # summary preferred over body
        self.assertNotIn(str(far.id), [s["filament_id"] for s in res.data["sources"]])

    def test_malformed_model_output_falls_back_to_single_uncited_segment(self):
        Filament.objects.create(
            type="text", title="t", body="x", status="done", embedding=sim_vec(0.9)
        )
        anthropic_client = MagicMock()
        anthropic_client.messages.create.return_value = SimpleNamespace(
            content=[SimpleNamespace(type="text", text="not json at all")]
        )
        openai_client = MagicMock()
        openai_client.embeddings.create.return_value = embedding_response(vec(1.0, 0.0))

        res = self._ask("anything", anthropic_client, openai_client)
        self.assertEqual(res.status_code, 200)  # never 500
        self.assertEqual(
            res.data["answer"], [{"text": "not json at all", "citation": None}]
        )
        self.assertEqual(res.data["sources"], [])
        self.assertEqual(res.data["follow_ups"], [])

    def test_empty_archive_answers_gracefully_without_claude(self):
        anthropic_client = MagicMock()
        openai_client = MagicMock()
        openai_client.embeddings.create.return_value = embedding_response(vec(1.0, 0.0))

        res = self._ask("anything", anthropic_client, openai_client)
        self.assertEqual(res.status_code, 200)
        self.assertIsNone(res.data["answer"][0]["citation"])
        self.assertEqual(res.data["sources"], [])
        anthropic_client.messages.create.assert_not_called()

    def test_blank_question_rejected(self):
        res = self.client.post("/api/v1/ask", {"question": "  "}, format="json")
        self.assertEqual(res.status_code, 400)


# ---------------------------------------------------------------------------
# Export (PRD v1 #12)
# ---------------------------------------------------------------------------


class ExportTests(AuthedAPITestCase):
    def _filament(self):
        f = Filament.objects.create(
            type="text", title="Deep Work", body="attention is finite",
            summary="On focus.", key_ideas=["focus is trainable"], status="done",
        )
        ActionItem.objects.create(filament=f, text="block two hours", done=True)
        f.tags.add(Tag.objects.create(name="deep-work"))
        other = Filament.objects.create(type="voice", title="Memo", body="y", status="done")
        FilamentLink.create_link(f, other, 0.91)
        return f, other

    def test_markdown_export(self):
        f, _ = self._filament()
        res = self.client.get(f"/api/v1/filaments/{f.id}/export?format=markdown")
        self.assertEqual(res.status_code, 200)
        self.assertIn("text/markdown", res["Content-Type"])
        self.assertIn('filename="deep-work.md"', res["Content-Disposition"])
        body = res.content.decode()
        for fragment in ("# Deep Work", "On focus.", "- focus is trainable",
                         "- [x] block two hours", "deep-work", "Memo (similarity 0.91)"):
            self.assertIn(fragment, body)

    def test_json_export_preserves_links(self):
        f, other = self._filament()
        res = self.client.get(f"/api/v1/filaments/{f.id}/export?format=json")
        self.assertEqual(res.status_code, 200)
        payload = json.loads(res.content)
        self.assertEqual(payload["tags"], ["deep-work"])
        self.assertEqual(payload["links"][0]["filament_id"], str(other.id))
        self.assertEqual(payload["action_items"], [{"text": "block two hours", "done": True}])

    def test_text_export(self):
        f, _ = self._filament()
        res = self.client.get(f"/api/v1/filaments/{f.id}/export?format=text")
        self.assertEqual(res.status_code, 200)
        self.assertIn("attention is finite", res.content.decode())

    def test_audio_export_returns_presigned_url_for_voice_only(self):
        f, voice = self._filament()
        voice.source_key = f"voice/{voice.id}.m4a"
        voice.save(update_fields=["source_key"])

        with patch("core.views.generate_download_url", return_value="https://signed"):
            res = self.client.get(f"/api/v1/filaments/{voice.id}/export?format=audio")
        self.assertEqual(res.status_code, 200)
        self.assertEqual(res.data["url"], "https://signed")

        res = self.client.get(f"/api/v1/filaments/{f.id}/export?format=audio")
        self.assertEqual(res.status_code, 400)  # text filament has no audio

    def test_unknown_format_rejected(self):
        f, _ = self._filament()
        res = self.client.get(f"/api/v1/filaments/{f.id}/export?format=docx")
        self.assertEqual(res.status_code, 400)


class HealthTests(APITestCase):
    def test_health_is_public(self):
        res = self.client.get("/health")
        self.assertEqual(res.status_code, 200)
        self.assertEqual(res.json(), {"status": "ok"})
